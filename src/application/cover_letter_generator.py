import os
import logging
import openai
from typing import Dict, Any
import re

logger = logging.getLogger(__name__)

class CoverLetterGenerator:
    """
    Generates customized cover letters based on job description using AI.
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the cover letter generator.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config
        self.user_info = config.get('USER_INFO', {})
        self.ai_settings = config.get('AI_SETTINGS', {})
        self.api_key = self.ai_settings.get('api_key', '')
        self.model = self.ai_settings.get('model', 'gpt-4')
        
        # Set up OpenAI API key
        if self.api_key:
            openai.api_key = self.api_key
        
    def generate(self, job: Dict[str, Any]) -> str:
        """
        Generate a cover letter based on job description.
        
        Args:
            job: Job listing dictionary
            
        Returns:
            Path to generated cover letter
        """
        if not self.api_key:
            logger.error("OpenAI API key not provided")
            return ""
        
        try:
            # Extract job information
            job_description = job.get('description', '')
            job_title = job.get('title', '')
            company_name = job.get('company_name', '')
            h1b_sponsor = job.get('sponsors_h1b', False)
            
            if not job_description or not job_title or not company_name:
                logger.error("Missing job information")
                return ""
            
            # Generate cover letter content
            cover_letter_content = self._generate_cover_letter_content(
                job_description, job_title, company_name, h1b_sponsor
            )
            
            if not cover_letter_content:
                logger.error("Failed to generate cover letter content")
                return ""
            
            # Create cover letter file
            output_path = self._create_cover_letter_file(cover_letter_content, job)
            
            logger.info(f"Created AI-powered cover letter: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error generating cover letter: {e}")
            return ""
    
    def _generate_cover_letter_content(self, job_description: str, job_title: str, 
                                      company_name: str, h1b_sponsor: bool = False) -> str:
        """
        Generate cover letter content using OpenAI.
        
        Args:
            job_description: Job description
            job_title: Job title
            company_name: Company name
            h1b_sponsor: Whether the company sponsors H1B visas
            
        Returns:
            Cover letter content
        """
        try:
            # Get user information
            user_name = self.user_info.get('name', 'Your Name')
            user_email = self.user_info.get('email', 'your.email@example.com')
            user_phone = self.user_info.get('phone', '123-456-7890')
            user_linkedin = self.user_info.get('linkedin', '')
            user_github = self.user_info.get('github', '')
            user_portfolio = self.user_info.get('portfolio', '')
            
            # Extract key skills and requirements from job description
            key_requirements = self._extract_key_requirements(job_description)
            
            # Add H1B sponsorship information if applicable
            visa_info = ""
            if h1b_sponsor:
                visa_info = "I've researched that your company sponsors H1B visas, which aligns with my long-term career goals."
            
            prompt = f"""
            You are an expert cover letter writer. Your task is to write a compelling, personalized cover letter for the following job:
            
            Job Title: {job_title}
            Company: {company_name}
            Job Description:
            {job_description}
            
            Key Requirements:
            {key_requirements}
            
            Applicant Information:
            Name: {user_name}
            Email: {user_email}
            Phone: {user_phone}
            LinkedIn: {user_linkedin}
            GitHub: {user_github}
            Portfolio: {user_portfolio}
            
            Additional Information:
            {visa_info}
            
            Write a professional cover letter that:
            1. Is addressed to the hiring manager at {company_name}
            2. Expresses genuine interest in the {job_title} position and the company
            3. Highlights relevant skills and experiences that match the job requirements
            4. Demonstrates knowledge of the company and industry
            5. Explains why the applicant is a good fit for the role
            6. Includes a call to action at the end
            
            The cover letter should be concise (300-400 words), professional in tone, and highly personalized.
            Format the letter properly with date, address, greeting, body paragraphs, closing, and signature.
            
            Return only the cover letter text.
            """
            
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert cover letter writer who creates highly personalized, compelling cover letters."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            cover_letter_content = response.choices[0].message.content.strip()
            return cover_letter_content
        
        except Exception as e:
            logger.error(f"Error generating cover letter content: {e}")
            return ""
    
    def _extract_key_requirements(self, job_description: str) -> str:
        """
        Extract key requirements from job description using AI.
        
        Args:
            job_description: Job description
            
        Returns:
            Key requirements as a bulleted list
        """
        try:
            prompt = f"""
            Extract the key skills, qualifications, and requirements from the following job description.
            Format the output as a bulleted list of the 5-8 most important requirements.
            
            Job Description:
            {job_description}
            """
            
            response = openai.ChatCompletion.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "You are an expert at analyzing job descriptions."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.3
            )
            
            key_requirements = response.choices[0].message.content.strip()
            return key_requirements
        
        except Exception as e:
            logger.error(f"Error extracting key requirements: {e}")
            return ""
    
    def _create_cover_letter_file(self, content: str, job: Dict[str, Any]) -> str:
        """
        Create cover letter file.
        
        Args:
            content: Cover letter content
            job: Job listing dictionary
            
        Returns:
            Path to cover letter file
        """
        try:
            # Create output directory if it doesn't exist
            output_dir = "data_folder/cover_letters"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate safe filename
            company_name = re.sub(r'[^\w\s-]', '', job.get('company_name', 'company')).strip()
            job_title = re.sub(r'[^\w\s-]', '', job.get('title', 'job')).strip()
            safe_filename = f"{company_name}_{job_title}_cover_letter".replace(' ', '_')
            
            # Create text file
            output_path = os.path.join(output_dir, f"{safe_filename}.txt")
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(content)
            
            # Also create a docx file for better formatting
            try:
                from docx import Document
                from docx.shared import Pt
                
                doc = Document()
                
                # Add date
                date_para = doc.add_paragraph()
                date_para.add_run(f"{time.strftime('%B %d, %Y')}").bold = True
                
                # Add spacing
                doc.add_paragraph()
                
                # Add greeting and content
                paragraphs = content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                
                # Save document
                docx_path = os.path.join(output_dir, f"{safe_filename}.docx")
                doc.save(docx_path)
                
                return docx_path
            
            except Exception as e:
                logger.warning(f"Could not create DOCX cover letter: {e}")
                return output_path
        
        except Exception as e:
            logger.error(f"Error creating cover letter file: {e}")
            return ""

