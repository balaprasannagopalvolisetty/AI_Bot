import os
import logging
import openai
from typing import Dict, Any
import PyPDF2
from docx import Document
from docx.shared import Pt
import re

logger = logging.getLogger(__name__)

class ResumeCustomizer:
    """
    Customizes resume based on job description, focusing only on skills and projects.
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        Initialize the resume customizer.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config
        self.user_info = config.get('USER_INFO', {})
        self.ai_settings = config.get('AI_SETTINGS', {})
        self.resume_path = self.user_info.get('resume_path', '')
        self.api_key = self.ai_settings.get('api_key', '')
        self.model = self.ai_settings.get('model', 'gpt-4')
        
        # Set up OpenAI API key
        if self.api_key:
            openai.api_key = self.api_key
        
    def customize(self, job: Dict[str, Any]) -> str:
        """
        Customize resume based on job description, focusing only on skills and projects.
        
        Args:
            job: Job listing dictionary
            
        Returns:
            Path to customized resume
        """
        if not self.resume_path or not os.path.exists(self.resume_path):
            logger.error(f"Resume file not found: {self.resume_path}")
            return self.resume_path
        
        if not self.api_key:
            logger.error("OpenAI API key not provided")
            return self.resume_path
        
        try:
            # Extract resume content
            resume_content = self._extract_resume_content()
            if not resume_content:
                logger.error("Failed to extract resume content")
                return self.resume_path
            
            # Extract job description
            job_description = job.get('description', '')
            job_title = job.get('title', '')
            company_name = job.get('company_name', '')
            
            if not job_description:
                logger.error("Job description is empty")
                return self.resume_path
            
            # Parse resume into sections
            resume_sections = self._parse_resume_sections(resume_content)
            
            # Generate customized skills and projects sections
            customized_sections = self._customize_targeted_sections(
                resume_sections, job_description, job_title, company_name
            )
            
            if not customized_sections:
                logger.error("Failed to generate customized resume sections")
                return self.resume_path
            
            # Merge customized sections with original resume
            merged_content = self._merge_resume_sections(resume_sections, customized_sections)
            
            # Create customized resume file
            output_path = self._create_customized_resume(merged_content, job)
            
            logger.info(f"Created customized resume with targeted skills and projects: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Error customizing resume: {e}")
            return self.resume_path
    
    def _extract_resume_content(self) -> str:
        """
        Extract text content from resume file.
        
        Returns:
            Resume text content
        """
        try:
            file_ext = os.path.splitext(self.resume_path)[1].lower()
            
            if file_ext == '.pdf':
                # Extract text from PDF
                with open(self.resume_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text()
                    return content
            
            elif file_ext == '.docx':
                # Extract text from DOCX
                doc = Document(self.resume_path)
                content = ""
                for para in doc.paragraphs:
                    content += para.text + "\n"
                return content
            
            elif file_ext == '.txt':
                # Extract text from TXT
                with open(self.resume_path, 'r', encoding='utf-8') as file:
                    return file.read()
            
            else:
                logger.error(f"Unsupported resume file format: {file_ext}")
                return ""
        
        except Exception as e:
            logger.error(f"Error extracting resume content: {e}")
            return ""
    
    def _parse_resume_sections(self, resume_content: str) -> Dict[str, str]:
        """
        Parse resume content into sections.
        
        Args:
            resume_content: Resume content
            
        Returns:
            Dictionary of resume sections
        """
        # Common section headers in resumes
        section_patterns = [
            (r'(?i)skills?|technical skills?|core competenc(?:y|ies)', 'skills'),
            (r'(?i)projects?|personal projects?|key projects?', 'projects'),
            (r'(?i)experience|work experience|professional experience|employment', 'experience'),
            (r'(?i)education|academic|qualifications', 'education'),
            (r'(?i)certifications?|licenses?', 'certifications'),
            (r'(?i)summary|profile|objective', 'summary'),
            (r'(?i)contact|personal information', 'contact')
        ]
        
        # Find section boundaries
        sections = {}
        section_starts = []
        
        # Find all potential section headers
        for pattern, section_name in section_patterns:
            for match in re.finditer(pattern, resume_content, re.MULTILINE):
                line_start = resume_content.rfind('\n', 0, match.start()) + 1
                line_end = resume_content.find('\n', match.end())
                if line_end == -1:
                    line_end = len(resume_content)
                
                # Get the full line containing the match
                line = resume_content[line_start:line_end].strip()
                
                # Check if this is likely a section header (short line, possibly with formatting)
                if len(line) < 50:  # Arbitrary threshold for header length
                    section_starts.append((line_start, section_name, line))
        
        # Sort section starts by position
        section_starts.sort()
        
        # Extract section content
        for i, (start, name, header) in enumerate(section_starts):
            end = section_starts[i+1][0] if i < len(section_starts) - 1 else len(resume_content)
            content = resume_content[start:end].strip()
            sections[name] = content
        
        # If no sections were found, use the entire resume
        if not sections:
            sections['full_resume'] = resume_content
        
        return sections
    
    def _customize_targeted_sections(self, resume_sections: Dict[str, str], 
                                    job_description: str, job_title: str, 
                                    company_name: str) -> Dict[str, str]:
        """
        Customize only the skills and projects sections of the resume.
        
        Args:
            resume_sections: Dictionary of resume sections
            job_description: Job description
            job_title: Job title
            company_name: Company name
            
        Returns:
            Dictionary of customized sections
        """
        customized_sections = {}
        
        try:
            # Only customize skills and projects sections
            sections_to_customize = ['skills', 'projects']
            
            for section_name in sections_to_customize:
                if section_name in resume_sections:
                    original_section = resume_sections[section_name]
                    
                    prompt = f"""
                    You are an expert resume customizer. Your task is to customize ONLY the {section_name} section of a resume to better match a job description.
                    
                    Original {section_name} section:
                    {original_section}
                    
                    Job Title: {job_title}
                    Company: {company_name}
                    Job Description:
                    {job_description}
                    
                    Please customize ONLY the {section_name} section to highlight relevant {section_name} that match the job description.
                    Keep the same format and structure as the original section, but modify the content to better match the job requirements.
                    Do not add fictional {section_name} that are not mentioned in the original section.
                    
                    Return only the customized {section_name} section text, starting with the section header.
                    """
                    
                    response = openai.ChatCompletion.create(
                        model=self.model,
                        messages=[
                            {"role": "system", "content": f"You are an expert resume customizer focusing on {section_name}."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=2000,
                        temperature=0.5
                    )
                    
                    customized_section = response.choices[0].message.content.strip()
                    customized_sections[section_name] = customized_section
            
            return customized_sections
        
        except Exception as e:
            logger.error(f"Error customizing targeted sections: {e}")
            return {}
    
    def _merge_resume_sections(self, original_sections: Dict[str, str], 
                              customized_sections: Dict[str, str]) -> str:
        """
        Merge customized sections with original resume.
        
        Args:
            original_sections: Dictionary of original resume sections
            customized_sections: Dictionary of customized sections
            
        Returns:
            Merged resume content
        """
        # If we couldn't parse sections properly, return the original full resume
        if 'full_resume' in original_sections:
            return original_sections['full_resume']
        
        # Start with a copy of the original sections
        merged_sections = original_sections.copy()
        
        # Replace only the customized sections
        for section_name, content in customized_sections.items():
            if section_name in merged_sections:
                merged_sections[section_name] = content
        
        # Reconstruct the resume in the original section order
        section_order = list(original_sections.keys())
        merged_content = ""
        
        for section_name in section_order:
            merged_content += merged_sections[section_name] + "\n\n"
        
        return merged_content.strip()
    
    def _create_customized_resume(self, content: str, job: Dict[str, Any]) -> str:
        """
        Create customized resume file.
        
        Args:
            content: Customized resume content
            job: Job listing dictionary
            
        Returns:
            Path to customized resume file
        """
        try:
            # Create output directory if it doesn't exist
            output_dir = "data_folder/customized_resumes"
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate safe filename
            company_name = re.sub(r'[^\w\s-]', '', job.get('company_name', 'company')).strip()
            job_title = re.sub(r'[^\w\s-]', '', job.get('title', 'job')).strip()
            safe_filename = f"{company_name}_{job_title}".replace(' ', '_')
            
            # Determine output file format based on original resume
            file_ext = os.path.splitext(self.resume_path)[1].lower()
            output_path = os.path.join(output_dir, f"{safe_filename}{file_ext}")
            
            if file_ext == '.pdf':
                # Create PDF (simplified - in reality, you'd need a PDF generation library)
                with open(output_path.replace('.pdf', '.txt'), 'w', encoding='utf-8') as file:
                    file.write(content)
                logger.warning("PDF generation not implemented. Created text file instead.")
                return output_path.replace('.pdf', '.txt')
            
            elif file_ext == '.docx':
                # Create DOCX
                doc = Document()
                for paragraph in content.split('\n'):
                    p = doc.add_paragraph(paragraph)
                    p.style.font.size = Pt(11)
                doc.save(output_path)
                return output_path
            
            else:
                # Create TXT
                output_path = os.path.join(output_dir, f"{safe_filename}.txt")
                with open(output_path, 'w', encoding='utf-8') as file:
                    file.write(content)
                return output_path
        
        except Exception as e:
            logger.error(f"Error creating customized resume: {e}")
            return self.resume_path

